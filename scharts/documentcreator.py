from docx import Document
from docx.shared import Inches
import yfinance as yf

def createdocument(stock_data,keys):

    document = Document()

    document.add_heading(stock_data['longName'], 0)

    p = document.add_paragraph(stock_data['longBusinessSummary'])
    #p.add_run('bold').bold = True
    #p.add_run(' and some ')
    #p.add_run('italic.').italic = True

    #document.add_heading('Heading, level 1', level=1)
    #document.add_paragraph('Intense quote', style='Intense Quote')

    #document.add_paragraph(
    #    'first item in unordered list', style='List Bullet'
    #)
    #document.add_paragraph(
    #    'first item in ordered list', style='List Number'
    #)

    #document.add_picture('CodecombatLogo.png', width=Inches(1.25))

    table = document.add_table(rows=1, cols=2)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'
    for key in keys:
        print(key,stock_data[key])
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(stock_data[key])

    document.add_page_break()

    document.save('demo.docx')


ticker = yf.Ticker('TSLA')
data = ticker.info
keys = ['currentPrice'
,'earningsGrowth'
,'currentRatio','returnOnAssets'
]
createdocument(data,keys)